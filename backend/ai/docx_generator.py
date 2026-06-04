import os
import json
import tempfile
import subprocess
from pathlib import Path
from datetime import datetime
import structlog
from utils.config import OUTPUTS_DIR

logger = structlog.get_logger(__name__)

NODE_SCRIPT = Path(__file__).parent / "generate_docx.js"

def create_resume_docx(profile_data: dict, tailored_data: dict, company_name: str) -> str:
    """
    Generates an ATS-friendly DOCX resume by calling a Node.js script.
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = profile_data.get('full_name', 'Candidate').replace(' ', '_')
    safe_company = company_name.replace(' ', '_')
    filename = f"{safe_name}_{safe_company}_{timestamp}.docx"
    
    filepath = OUTPUTS_DIR / filename
    
    data = {
        "profile_dict": profile_data,
        "tailored_data": tailored_data,
        "company_name": company_name
    }
    
    with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.json') as tmp:
        json.dump(data, tmp)
        tmp_path = tmp.name
        
    try:
        # Call Node.js script
        result = subprocess.run(
            ["node", str(NODE_SCRIPT), tmp_path, str(filepath)],
            capture_output=True, text=True, check=True
        )
        logger.info("node_docx_generated", output=result.stdout)
    except subprocess.CalledProcessError as e:
        logger.error("node_docx_generation_failed", error=e.stderr)
        raise Exception(f"DOCX Generation failed: {e.stderr}")
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
            
    return str(filepath)

def create_cover_letter_docx(profile_data: dict, content: str, company_name: str) -> str:
    """
    Placeholder: The Node.js script currently generates Resumes.
    We'll leave this unimplemented for now, or just return empty.
    Wait, the user said we need to rewrite `docx_generator.py`. 
    Actually, let's just make a simple text-based DOCX using python-docx for the cover letter since it's just text.
    But the node script can also be extended.
    To avoid breaking cover letter, I'll keep the python-docx logic just for cover letter.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(profile_data.get('full_name', 'Full Name'))
    name_run.bold = True
    name_run.font.size = Pt(14)
    
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_info = []
    if profile_data.get('email'): contact_info.append(profile_data['email'])
    if profile_data.get('phone'): contact_info.append(profile_data['phone'])
    if profile_data.get('location'): contact_info.append(profile_data['location'])
        
    contact_run = contact_p.add_run(" | ".join(contact_info))
    contact_run.font.size = Pt(11)
    
    doc.add_paragraph() # spacing
    doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph() # spacing
    
    for paragraph in content.split('\n\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = profile_data.get('full_name', 'Candidate').replace(' ', '_')
    safe_company = company_name.replace(' ', '_')
    filename = f"CL_{safe_name}_{safe_company}_{timestamp}.docx"
    
    filepath = OUTPUTS_DIR / filename
    doc.save(str(filepath))
    
    return str(filepath)
